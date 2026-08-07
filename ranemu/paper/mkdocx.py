#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ranemu.paper.mkdocx — 논문 Markdown → 투고용 Word(.docx) 변환기.

왜 직접 만드는가
================
pandoc/libreoffice 가 이 환경에 없다. 그리고 어차피 학술지 투고본은 일반
Markdown 변환으로는 부족하다 — 연속 줄번호(MDPI 필수), 표/그림 번호와 캡션
위치 규약(표 캡션은 위, 그림 캡션은 아래), 초록/키워드 블록, 참고문헌 스타일이
저널마다 다르다. 그래서 필요한 만큼만 정확히 지원하는 변환기를 둔다.

지원하는 Markdown 부분집합
--------------------------
    ---            YAML 프런트매터(title/authors/affiliations/abstract/keywords/…)
    # / ## / ###   제목 계층
    문단           **굵게** *기울임* `고정폭` 지원
    | a | b |      GFM 파이프 표 (바로 앞 줄의 `**Table N.** …` 를 캡션으로 사용)
    ![cap](path)   그림 (캡션은 그림 아래)
    - / 1.         목록
    $$ ... $$      별행 수식(가운데 정렬)
    > ...          주석 블록
    [1] ...        참고문헌(마지막 References 절에서 자동 들여쓰기)

사용:
    python3 -m ranemu.paper.mkdocx paper.md -o paper.docx --style mdpi
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# 저널별 서식
# ─────────────────────────────────────────────────────────────────────────────
STYLES: Dict[str, Dict[str, Any]] = {
    # Elsevier(ICT Express) — 투고본은 1단·이중행간 권장, 줄번호 권장
    "elsevier": {
        "font": "Times New Roman", "size": 12, "line_spacing": 2.0,
        "margins_cm": 2.5, "line_numbers": True, "title_size": 16,
        "abstract_indent_cm": 0.0, "caption_size": 10,
        "table_font_size": 9, "ref_hanging": True,
    },
    # MDPI(Sensors/Electronics) — 투고본 1단, 연속 줄번호 필수
    "mdpi": {
        "font": "Palatino Linotype", "size": 10, "line_spacing": 1.15,
        "margins_cm": 2.0, "line_numbers": True, "title_size": 14,
        "abstract_indent_cm": 0.0, "caption_size": 9,
        "table_font_size": 8, "ref_hanging": False,
    },
    "plain": {
        "font": "Calibri", "size": 11, "line_spacing": 1.5,
        "margins_cm": 2.5, "line_numbers": False, "title_size": 16,
        "abstract_indent_cm": 0.0, "caption_size": 10,
        "table_font_size": 9, "ref_hanging": True,
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# 프런트매터
# ─────────────────────────────────────────────────────────────────────────────
def parse_front_matter(text: str) -> Tuple[Dict[str, Any], str]:
    """--- 로 감싼 간이 YAML 을 읽는다(중첩 없음, 리스트는 '- ' 항목)."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end < 0:
        return {}, text
    raw = text[3:end].strip("\n")
    body = text[end + 4:].lstrip("\n")
    meta: Dict[str, Any] = {}
    key: Optional[str] = None
    for line in raw.split("\n"):
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        if line.startswith("  - ") or line.startswith("- "):
            if key:
                meta.setdefault(key, [])
                if isinstance(meta[key], list):
                    meta[key].append(line.split("- ", 1)[1].strip().strip('"'))
            continue
        if ":" in line:
            k, v = line.split(":", 1)
            key = k.strip()
            v = v.strip().strip('"')
            meta[key] = v if v else []
    return meta, body


# ─────────────────────────────────────────────────────────────────────────────
# docx 저수준 도우미
# ─────────────────────────────────────────────────────────────────────────────
def _enable_line_numbers(section) -> None:
    """연속 줄번호(MDPI 요구사항)."""
    sect_pr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sect_pr.append(ln)


def _add_page_numbers(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr in ("begin", "text", "separate", "end"):
        el = OxmlElement("w:fldChar") if instr != "text" else OxmlElement("w:instrText")
        if instr == "text":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)


def _set_cell_bg(cell, hexcolor: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_borders(cell, top: bool = False, bottom: bool = False,
                  size: int = 8) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, on in (("top", top), ("bottom", bottom)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single" if on else "nil")
        if on:
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


# ─────────────────────────────────────────────────────────────────────────────
# 인라인 서식
# ─────────────────────────────────────────────────────────────────────────────
_INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")


def add_runs(par, text: str, base_font: str, size: float) -> None:
    """**굵게** *기울임* `고정폭` 을 run 으로 분해해 붙인다."""
    text = text.replace("\\*", "")
    for tok in _INLINE.split(text):
        if not tok:
            continue
        tok = tok.replace("", "*")
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
        elif (tok.startswith("*") and tok.endswith("*") and len(tok) > 2
              and not tok.startswith("**")):
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            r = par.add_run(tok)
        if r.font.name is None:
            r.font.name = base_font
        if r.font.size is None:
            r.font.size = Pt(size)


# ─────────────────────────────────────────────────────────────────────────────
# 변환기
# ─────────────────────────────────────────────────────────────────────────────
class DocxBuilder:
    def __init__(self, style: str = "plain", base_dir: str = "."):
        self.cfg = STYLES.get(style, STYLES["plain"])
        self.style_name = style
        self.base_dir = base_dir
        self.doc = Document()
        self._setup()
        self.in_references = False

    # ── 문서 기본 ─────────────────────────────────────────────────────────
    def _setup(self) -> None:
        cfg = self.cfg
        sec = self.doc.sections[0]
        m = Cm(cfg["margins_cm"])
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m
        if cfg["line_numbers"]:
            _enable_line_numbers(sec)
        _add_page_numbers(sec)

        st = self.doc.styles["Normal"]
        st.font.name = cfg["font"]
        st.font.size = Pt(cfg["size"])
        st._element.rPr.rFonts.set(qn("w:eastAsia"), cfg["font"])
        pf = st.paragraph_format
        pf.line_spacing = cfg["line_spacing"]
        pf.space_after = Pt(0 if cfg["line_spacing"] >= 1.5 else 6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _p(self, text: str = "", *, align=None, size=None, bold=False,
           italic=False, space_before=0, space_after=0, indent_cm=None,
           spacing=None, color=None):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        if indent_cm is not None:
            pf.left_indent = Cm(indent_cm)
        if spacing is not None:
            pf.line_spacing = spacing
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.name = self.cfg["font"]
            r.font.size = Pt(size or self.cfg["size"])
            if color:
                r.font.color.rgb = RGBColor.from_string(color)
        return p

    # ── 표제부 ────────────────────────────────────────────────────────────
    def front_matter(self, meta: Dict[str, Any]) -> None:
        cfg = self.cfg
        if meta.get("title"):
            self._p(meta["title"], align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=cfg["title_size"], bold=True, space_after=10,
                    spacing=1.0)
        authors = meta.get("authors") or []
        if authors:
            self._p("; ".join(authors) if isinstance(authors, list) else str(authors),
                    align=WD_ALIGN_PARAGRAPH.CENTER, size=cfg["size"],
                    space_after=4, spacing=1.0)
        affs = meta.get("affiliations") or []
        if isinstance(affs, list):
            for i, a in enumerate(affs, 1):
                self._p(f"{i} {a}", align=WD_ALIGN_PARAGRAPH.CENTER,
                        size=cfg["size"] - 1, italic=True, space_after=2,
                        spacing=1.0)
        if meta.get("corresponding"):
            self._p(f"* Correspondence: {meta['corresponding']}",
                    align=WD_ALIGN_PARAGRAPH.CENTER, size=cfg["size"] - 1,
                    space_after=10, spacing=1.0)

        if meta.get("abstract"):
            p = self._p(space_before=8, space_after=6,
                        indent_cm=cfg["abstract_indent_cm"], spacing=1.15)
            r = p.add_run("Abstract: ")
            r.bold = True
            r.font.name = cfg["font"]; r.font.size = Pt(cfg["size"])
            add_runs(p, meta["abstract"], cfg["font"], cfg["size"])
        if meta.get("keywords"):
            kw = meta["keywords"]
            kw = "; ".join(kw) if isinstance(kw, list) else str(kw)
            p = self._p(space_after=12, indent_cm=cfg["abstract_indent_cm"],
                        spacing=1.15)
            r = p.add_run("Keywords: ")
            r.bold = True
            r.font.name = cfg["font"]; r.font.size = Pt(cfg["size"])
            add_runs(p, kw, cfg["font"], cfg["size"])

    # ── 블록 요소 ─────────────────────────────────────────────────────────
    def heading(self, level: int, text: str) -> None:
        sizes = {1: self.cfg["size"] + 3, 2: self.cfg["size"] + 1,
                 3: self.cfg["size"]}
        self.in_references = bool(re.match(r"^\s*(\d+\.\s*)?references\s*$",
                                           text, re.I))
        self._p(text, size=sizes.get(level, self.cfg["size"]), bold=True,
                space_before=10 if level <= 2 else 8, space_after=4,
                align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0)

    def paragraph(self, text: str) -> None:
        cfg = self.cfg
        if self.in_references and re.match(r"^\[\d+\]", text):
            p = self._p(space_after=2, spacing=1.0)
            if cfg["ref_hanging"]:
                p.paragraph_format.left_indent = Cm(0.9)
                p.paragraph_format.first_line_indent = Cm(-0.9)
            add_runs(p, text, cfg["font"], cfg["size"] - 1)
            return
        p = self._p(space_after=4 if cfg["line_spacing"] < 1.5 else 0)
        add_runs(p, text, cfg["font"], cfg["size"])

    def note(self, text: str) -> None:
        p = self._p(indent_cm=0.8, space_before=4, space_after=4, spacing=1.0)
        add_runs(p, text, self.cfg["font"], self.cfg["size"] - 1)

    def bullet(self, text: str, numbered: bool = False) -> None:
        p = self.doc.add_paragraph(style="List Number" if numbered
                                   else "List Bullet")
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, text, self.cfg["font"], self.cfg["size"])

    def equation(self, text: str) -> None:
        p = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6,
                    space_after=6, spacing=1.0)
        r = p.add_run(text)
        r.italic = True
        r.font.name = "Cambria Math"
        r.font.size = Pt(self.cfg["size"])

    def caption(self, text: str, above: bool = True) -> None:
        p = self._p(space_before=6 if above else 3,
                    space_after=3 if above else 8, spacing=1.0,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
        add_runs(p, text, self.cfg["font"], self.cfg["caption_size"])

    def table(self, rows: List[List[str]]) -> None:
        if not rows:
            return
        cfg = self.cfg
        ncol = max(len(r) for r in rows)
        t = self.doc.add_table(rows=len(rows), cols=ncol)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for i, row in enumerate(rows):
            for j in range(ncol):
                cell = t.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                val = row[j] if j < len(row) else ""
                # 숫자 열은 오른쪽 정렬(읽기 편하고 학술지 관행)
                if re.fullmatch(r"[−\-+]?[\d.,]+(\s*±\s*[\d.]+)?%?", val.strip()):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_runs(p, val, cfg["font"], cfg["table_font_size"])
                if i == 0:
                    for r in p.runs:
                        r.bold = True
                    _set_cell_bg(cell, "F2F2F2")
                _cell_borders(cell, top=(i == 0),
                              bottom=(i == 0 or i == len(rows) - 1))
        self._p(space_after=2, spacing=1.0)

    def figure(self, path: str, caption: str) -> None:
        full = path if os.path.isabs(path) else os.path.join(self.base_dir, path)
        if not os.path.exists(full):
            for ext in (".png", ".pdf"):
                cand = os.path.splitext(full)[0] + ext
                if os.path.exists(cand):
                    full = cand
                    break
        p = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8,
                    space_after=2, spacing=1.0)
        if os.path.exists(full):
            sec = self.doc.sections[0]
            avail = sec.page_width - sec.left_margin - sec.right_margin
            p.add_run().add_picture(full, width=Emu(int(avail * 0.96)))
        else:
            r = p.add_run(f"[figure not found: {path}]")
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        if caption:
            self.caption(caption, above=False)

    def pagebreak(self) -> None:
        self.doc.add_page_break()

    def save(self, path: str) -> None:
        self.doc.save(path)


# ─────────────────────────────────────────────────────────────────────────────
# 파서
# ─────────────────────────────────────────────────────────────────────────────
_FIG = re.compile(r"^!\[(?P<cap>.*?)\]\((?P<path>[^)]+)\)\s*$")
_TBLROW = re.compile(r"^\s*\|.*\|\s*$")
_SEP = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def convert(md_path: str, out_path: str, style: str = "plain",
            base_dir: Optional[str] = None) -> None:
    with open(md_path, encoding="utf-8") as f:
        text = f.read()
    meta, body = parse_front_matter(text)
    b = DocxBuilder(style=style,
                    base_dir=base_dir or os.path.dirname(os.path.abspath(md_path)))
    b.front_matter(meta)

    lines = body.split("\n")
    i = 0
    para: List[str] = []
    pending_caption: Optional[str] = None

    def flush() -> None:
        nonlocal para
        if para:
            b.paragraph(" ".join(x.strip() for x in para))
            para = []

    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if not s:
            flush()
            i += 1
            continue

        if s.startswith("<!--"):                       # 주석
            flush()
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if s in ("---", "***", "___"):                 # 구분선 → 무시
            flush(); i += 1; continue

        if s == "\\pagebreak":
            flush(); b.pagebreak(); i += 1; continue

        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            flush()
            b.heading(len(m.group(1)), m.group(2).strip())
            i += 1
            continue

        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            flush(); b.equation(s[2:-2].strip()); i += 1; continue
        if s == "$$":
            flush()
            i += 1
            buf = []
            while i < len(lines) and lines[i].strip() != "$$":
                buf.append(lines[i].strip()); i += 1
            i += 1
            b.equation("  ".join(buf))
            continue

        mf = _FIG.match(s)
        if mf:
            flush()
            b.figure(mf.group("path"), mf.group("cap"))
            pending_caption = None
            i += 1
            continue

        if s.startswith("> "):
            flush()
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip()); i += 1
            b.note(" ".join(buf))
            continue

        mb = re.match(r"^[-*+]\s+(.*)$", s)
        mn = re.match(r"^\d+[.)]\s+(.*)$", s)
        if mb or mn:
            flush()
            b.bullet((mb or mn).group(1), numbered=bool(mn))
            i += 1
            continue

        # 표 캡션(표 바로 앞의 **Table N.** 로 시작하는 줄)
        if re.match(r"^\*\*Table\s", s) and i + 1 < len(lines) \
                and _TBLROW.match(lines[i + 1].strip()):
            flush()
            b.caption(s, above=True)
            i += 1
            continue

        if _TBLROW.match(s):
            flush()
            rows: List[List[str]] = []
            while i < len(lines) and _TBLROW.match(lines[i].strip()):
                raw = lines[i].strip()
                if not _SEP.match(raw):
                    cells = [c.strip() for c in raw.strip("|").split("|")]
                    rows.append(cells)
                i += 1
            b.table(rows)
            continue

        para.append(s)
        i += 1

    flush()
    b.save(out_path)
    print(f"작성: {out_path}  (style={style})")


def main(argv: Optional[List[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="논문 Markdown → Word")
    ap.add_argument("input")
    ap.add_argument("-o", "--output", default=None)
    ap.add_argument("--style", default="plain", choices=sorted(STYLES))
    ap.add_argument("--base-dir", default=None, help="그림 경로 기준")
    a = ap.parse_args(argv)
    out = a.output or os.path.splitext(a.input)[0] + ".docx"
    convert(a.input, out, style=a.style, base_dir=a.base_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
